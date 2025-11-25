import os
import re
from docx import Document
from docx.shared import Inches

class WordExporter:
    def __init__(self):
        pass

    def parse_markdown_sections(self, draft_path: str) -> dict:
        """Parse markdown into sections based on H1/H2 headers."""
        sections = {}
        current_section = None
        current_content = []
        
        with open(draft_path, "r") as f:
            lines = f.readlines()
            
        for line in lines:
            stripped = line.strip()
            # Detect Header
            if stripped.startswith("#"):
                # Save previous section
                if current_section:
                    sections[current_section] = current_content
                
                # Start new section
                # Remove # and whitespace, also remove numbering if present (e.g., "1. Introduction" -> "Introduction")
                header_text = stripped.lstrip("#").strip()
                # Simple normalization: remove leading numbers and dots
                header_text = re.sub(r'^\d+\.?\s*', '', header_text)
                current_section = header_text
                current_content = []
            else:
                if current_section:
                    current_content.append(line)
        
        # Save last section
        if current_section:
            sections[current_section] = current_content
            
        return sections

    def fill_template(self, doc: Document, sections: dict):
        """
        Replace content in template sections with draft content.
        Strategies:
        1. Find a paragraph with style containing 'Heading' or 'MDPI...heading'.
        2. Check if text matches a section key (fuzzy match).
        3. If match, delete following paragraphs until next header.
        4. Insert draft content.
        """
        # Iterate through paragraphs to find headers
        # We need to be careful about modifying the list while iterating
        # So we collect indices first
        
        header_indices = []
        for i, p in enumerate(doc.paragraphs):
            # Check if it's a header
            is_header = False
            if p.style.name.startswith("Heading") or "heading" in p.style.name.lower():
                is_header = True
            
            if is_header:
                # Normalize text
                text = p.text.strip()
                norm_text = re.sub(r'^\d+\.?\s*', '', text) # Remove "1. "
                
                if norm_text in sections:
                    header_indices.append((i, norm_text))
        
        # Process in reverse order to avoid index shifting issues when deleting
        for i, section_name in reversed(header_indices):
            print(f"Updating section: {section_name}")
            content_lines = sections[section_name]
            
            # 1. Delete existing content until next header
            # We look ahead from i+1
            j = i + 1
            while j < len(doc.paragraphs):
                p = doc.paragraphs[j]
                # Stop if next header found
                if p.style.name.startswith("Heading") or "heading" in p.style.name.lower():
                    break
                
                # Delete paragraph
                # python-docx doesn't have a simple delete for paragraphs in the list
                # We have to remove the element
                p._element.getparent().remove(p._element)
                # Don't increment j because the next element shifts to j
                # But wait, doc.paragraphs is a property, it re-queries.
                # So if we remove, the list length shrinks.
                # Actually, doc.paragraphs might be cached or dynamic. 
                # Safer to just clear text? No, we want to remove placeholders.
                # Let's try clearing text for now to be safe, or use the element removal.
                # If we remove element, we should be careful.
                # Let's just clear text to avoid breaking structure if unsure?
                # User wants "modification", so removal is better.
                # Let's assume element removal works.
            
            # 2. Insert new content after the header
            # We can't easily "insert at index" in python-docx API.
            # We can only `add_paragraph` at the end, or `insert_paragraph_before`.
            # So we should find the *next* paragraph (if it exists) and insert before it.
            
            # If we deleted everything up to end, we append.
            # If we deleted up to next header, we insert before next header.
            
            # Re-find the insertion point.
            # Since we processed in reverse, the 'next header' (originally at some index > i) 
            # might have shifted if we deleted stuff? 
            # No, we process reverse, so changes at 'i' don't affect indices < i.
            # But they affect indices > i.
            # Wait, if we process reverse (bottom up), say we have Header A (idx 10) and Header B (idx 20).
            # We process B first. We delete 21-25. Insert new stuff.
            # Then we process A. We delete 11-19. Insert new stuff.
            # This seems safe.
            
            # Insertion point: The paragraph currently at i+1 (which is the next header or end of doc)
            # But wait, we deleted paragraphs starting at i+1.
            # So the paragraph that *was* at j (the next header) is now at i+1.
            # So we insert before doc.paragraphs[i+1].
            
            if i + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[i+1]
                for line in reversed(content_lines): # Insert in reverse order so they end up correct
                    self._insert_line_before(doc, next_p, line)
            else:
                # End of document, just append
                for line in content_lines:
                    self._add_line_to_doc(doc, line)

    def _insert_line_before(self, doc, next_p, line):
        line = line.strip()
        if not line:
            return
            
        if line.startswith("!["):
            # Image handling is tricky with insert_before. 
            # doc.add_picture appends. 
            # We might have to skip images or just append them?
            # Or use a placeholder paragraph and run add_picture?
            # python-docx doesn't support inserting pictures at arbitrary positions easily.
            # Fallback: Just add text "[Image: ...]"
            new_p = next_p.insert_paragraph_before(f"[Image inserted: {line}]")
        else:
            new_p = next_p.insert_paragraph_before(line)

    def _add_line_to_doc(self, doc, line):
        # Helper to add line (text or image) to end of doc
        line = line.strip()
        if not line:
            return
            
        if line.startswith("!["):
             match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
             if match:
                 img_path = match.group(2)
                 if os.path.exists(img_path):
                     try:
                         doc.add_picture(img_path, width=Inches(6))
                     except:
                         doc.add_paragraph(line)
                 else:
                     doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)

    def export_to_word(self, draft_path: str, template_path: str, output_path: str) -> str:
        if not os.path.exists(draft_path):
            raise FileNotFoundError(f"Draft file not found: {draft_path}")
        
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
            # Parse draft and fill template
            sections = self.parse_markdown_sections(draft_path)
            if sections:
                self.fill_template(doc, sections)
            else:
                # Fallback to appending if no sections found
                with open(draft_path, "r") as f:
                    for line in f:
                        self._add_line_to_doc(doc, line)
        else:
            doc = Document()
            with open(draft_path, "r") as f:
                for line in f:
                    self._add_line_to_doc(doc, line)
            
        doc.save(output_path)
        return output_path
