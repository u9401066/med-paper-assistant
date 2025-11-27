"""
Word Writer - Low-level Word document manipulation.

This module provides precise control over Word document content,
allowing the MCP Agent (AI) to specify exactly where content should be placed.
The Agent makes decisions, this module executes them.
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass


@dataclass
class InsertInstruction:
    """Instructions for inserting content into a document."""
    section_name: str
    content: str
    position: str = "replace"  # "replace", "append", "prepend"
    style: Optional[str] = None


class WordWriter:
    """
    Executes content insertion into Word documents.
    
    This class does NOT make decisions about where content goes.
    The MCP Agent provides explicit instructions via InsertInstruction.
    """
    
    def __init__(self):
        pass
    
    def create_document_from_template(self, template_path: str) -> Document:
        """Load a template as a new document."""
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
        return Document(template_path)
    
    def create_blank_document(self) -> Document:
        """Create a new blank document."""
        return Document()
    
    def find_section_indices(self, doc: Document, section_name: str) -> Tuple[int, int]:
        """
        Find the start and end paragraph indices for a section.
        
        Args:
            doc: The Word document.
            section_name: Name of the section to find (case-insensitive, partial match).
            
        Returns:
            Tuple of (start_index, end_index) where:
            - start_index is the header paragraph index
            - end_index is the index of the next header (or end of doc)
        """
        section_lower = section_name.lower()
        start_idx = None
        end_idx = len(doc.paragraphs)
        
        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name.lower()
            is_heading = "heading" in style_name
            
            if is_heading:
                text_lower = para.text.lower()
                # Normalize: remove numbers like "1. "
                text_normalized = re.sub(r'^\d+\.?\s*', '', text_lower).strip()
                
                if start_idx is not None:
                    # We've found our section, now looking for the next header
                    end_idx = i
                    break
                elif section_lower in text_normalized or text_normalized in section_lower:
                    start_idx = i
        
        if start_idx is None:
            raise ValueError(f"Section '{section_name}' not found in document")
        
        return start_idx, end_idx
    
    def clear_section_content(self, doc: Document, section_name: str) -> int:
        """
        Clear all content from a section (keeping the header).
        
        Returns:
            Number of paragraphs removed.
        """
        start_idx, end_idx = self.find_section_indices(doc, section_name)
        
        removed = 0
        # Remove paragraphs from start_idx+1 to end_idx-1 (keep header and next header)
        # We need to remove in reverse to avoid index shifting issues
        paragraphs_to_remove = []
        for i in range(start_idx + 1, end_idx):
            paragraphs_to_remove.append(doc.paragraphs[i])
        
        for para in paragraphs_to_remove:
            para._element.getparent().remove(para._element)
            removed += 1
        
        return removed
    
    def insert_content_in_section(
        self, 
        doc: Document, 
        section_name: str, 
        content_paragraphs: List[str],
        clear_existing: bool = True
    ) -> int:
        """
        Insert content into a specific section.
        
        Args:
            doc: The Word document.
            section_name: Target section name.
            content_paragraphs: List of paragraph texts to insert.
            clear_existing: If True, clear existing content first.
            
        Returns:
            Number of paragraphs inserted.
        """
        if clear_existing:
            self.clear_section_content(doc, section_name)
        
        # Re-find indices after clearing (indices may have shifted)
        start_idx, end_idx = self.find_section_indices(doc, section_name)
        
        # After clearing, the next section header (if any) is now at start_idx + 1
        # We need to insert before it
        inserted = 0
        if start_idx + 1 < len(doc.paragraphs):
            insert_before = doc.paragraphs[start_idx + 1]
            # Insert in NORMAL order - insert_paragraph_before inserts directly 
            # before the target, preserving previous insertions
            for para_text in content_paragraphs:
                para_text = para_text.strip()
                if para_text:
                    insert_before.insert_paragraph_before(para_text)
                    inserted += 1
        else:
            # Section is at end of document, just append
            for para_text in content_paragraphs:
                para_text = para_text.strip()
                if para_text:
                    doc.add_paragraph(para_text)
                    inserted += 1
        
        return inserted
    
    def execute_instructions(
        self, 
        doc: Document, 
        instructions: List[InsertInstruction]
    ) -> Dict[str, int]:
        """
        Execute a batch of insert instructions.
        
        This is the main entry point for the MCP Agent to modify a document.
        The Agent provides a list of instructions, and this method executes them.
        
        Args:
            doc: The Word document.
            instructions: List of InsertInstruction objects.
            
        Returns:
            Dict mapping section names to number of paragraphs inserted.
        """
        results = {}
        
        for instr in instructions:
            try:
                # Split content into paragraphs
                paragraphs = [p for p in instr.content.split('\n') if p.strip()]
                
                if instr.position == "replace":
                    count = self.insert_content_in_section(
                        doc, instr.section_name, paragraphs, clear_existing=True
                    )
                elif instr.position == "append":
                    count = self.insert_content_in_section(
                        doc, instr.section_name, paragraphs, clear_existing=False
                    )
                else:
                    count = 0
                
                results[instr.section_name] = count
                
            except Exception as e:
                results[instr.section_name] = f"Error: {str(e)}"
        
        return results
    
    def save_document(self, doc: Document, output_path: str) -> str:
        """Save the document to disk."""
        # Ensure output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        doc.save(output_path)
        return output_path
    
    def get_document_text(self, doc: Document) -> str:
        """Extract all text from the document for verification."""
        return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    
    def count_words_in_section(self, doc: Document, section_name: str) -> int:
        """Count words in a specific section."""
        try:
            start_idx, end_idx = self.find_section_indices(doc, section_name)
        except ValueError:
            return 0
        
        word_count = 0
        for i in range(start_idx + 1, end_idx):
            text = doc.paragraphs[i].text
            # Simple word count
            words = [w for w in text.split() if w.strip()]
            word_count += len(words)
        
        return word_count
    
    def get_all_word_counts(self, doc: Document) -> Dict[str, int]:
        """Get word counts for all sections."""
        counts = {}
        current_section = None
        current_words = 0
        
        for para in doc.paragraphs:
            style_name = para.style.name.lower()
            is_heading = "heading" in style_name
            
            if is_heading and para.text.strip():
                # Save previous section
                if current_section:
                    counts[current_section] = current_words
                
                # Start new section
                section_name = re.sub(r'^\d+\.?\s*', '', para.text.strip())
                current_section = section_name
                current_words = 0
            elif current_section:
                words = [w for w in para.text.split() if w.strip()]
                current_words += len(words)
        
        # Save last section
        if current_section:
            counts[current_section] = current_words
        
        return counts
