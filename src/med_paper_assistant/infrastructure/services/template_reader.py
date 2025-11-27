"""
Template Reader - Provides template structure information for MCP Agent decision-making.

This module reads Word templates and extracts structural information,
allowing the MCP Agent (AI) to make intelligent decisions about content placement.
"""

import os
from docx import Document
from dataclasses import dataclass
from typing import List, Dict, Optional


@dataclass
class TemplateSection:
    """Represents a section in the template."""
    index: int
    name: str
    style: str
    placeholder_text: str
    word_limit: Optional[int] = None
    is_required: bool = True


@dataclass
class TemplateStructure:
    """Complete template structure."""
    name: str
    sections: List[TemplateSection]
    total_word_limit: Optional[int] = None
    journal_name: Optional[str] = None
    

class TemplateReader:
    """Reads and analyzes Word template structures."""
    
    # Common word limits for different sections (can be overridden by template metadata)
    DEFAULT_WORD_LIMITS = {
        "abstract": 250,
        "introduction": 800,
        "methods": 1500,
        "results": 1500,
        "discussion": 1500,
        "conclusions": 300,
    }
    
    def __init__(self, templates_dir: str = None):
        if templates_dir:
            self.templates_dir = templates_dir
        else:
            # Default: project_root/templates
            self.templates_dir = os.path.join(
                os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))),
                "templates"
            )
    
    def read_template(self, template_path: str) -> TemplateStructure:
        """
        Read a Word template and extract its structure.
        
        Args:
            template_path: Path to the .docx template file.
            
        Returns:
            TemplateStructure with all sections and their properties.
        """
        if not os.path.isabs(template_path):
            template_path = os.path.join(self.templates_dir, template_path)
            
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        doc = Document(template_path)
        sections = []
        
        for i, para in enumerate(doc.paragraphs):
            # Detect headers/sections
            style_name = para.style.name.lower()
            is_heading = (
                style_name.startswith("heading") or
                "heading" in style_name or
                "title" in style_name
            )
            
            if is_heading and para.text.strip():
                # Extract section name (remove numbering like "1. ")
                import re
                section_name = re.sub(r'^\d+\.?\s*', '', para.text.strip())
                
                # Estimate word limit based on section type
                word_limit = None
                section_lower = section_name.lower()
                for key, limit in self.DEFAULT_WORD_LIMITS.items():
                    if key in section_lower:
                        word_limit = limit
                        break
                
                sections.append(TemplateSection(
                    index=i,
                    name=section_name,
                    style=para.style.name,
                    placeholder_text=para.text.strip(),
                    word_limit=word_limit,
                    is_required=self._is_required_section(section_name)
                ))
        
        return TemplateStructure(
            name=os.path.basename(template_path),
            sections=sections,
            journal_name=self._extract_journal_name(doc)
        )
    
    def _is_required_section(self, section_name: str) -> bool:
        """Determine if a section is required."""
        required = ["introduction", "methods", "results", "discussion", "abstract"]
        optional = ["acknowledgments", "appendix", "supplementary", "patents", "abbreviations"]
        
        section_lower = section_name.lower()
        for opt in optional:
            if opt in section_lower:
                return False
        return True
    
    def _extract_journal_name(self, doc: Document) -> Optional[str]:
        """Try to extract journal name from template metadata or content."""
        # Check first few paragraphs for journal info
        for para in doc.paragraphs[:5]:
            text = para.text.lower()
            if "sensors" in text:
                return "Sensors (MDPI)"
            elif "journal" in text:
                return para.text.strip()
        return None
    
    def get_template_summary(self, template_path: str) -> str:
        """
        Get a human-readable summary of template structure.
        This is what the MCP Agent will use to make decisions.
        
        Returns:
            Markdown-formatted summary string.
        """
        structure = self.read_template(template_path)
        
        output = f"# Template: {structure.name}\n\n"
        if structure.journal_name:
            output += f"**Journal**: {structure.journal_name}\n\n"
        
        output += "## Sections\n\n"
        output += "| # | Section | Style | Word Limit | Required |\n"
        output += "|---|---------|-------|------------|----------|\n"
        
        for i, section in enumerate(structure.sections, 1):
            limit_str = str(section.word_limit) if section.word_limit else "-"
            req_str = "✓" if section.is_required else "○"
            output += f"| {i} | {section.name} | {section.style} | {limit_str} | {req_str} |\n"
        
        output += "\n## Usage Notes\n"
        output += "- ✓ = Required section\n"
        output += "- ○ = Optional section\n"
        output += "- Word limits are guidelines based on typical journal requirements\n"
        
        return output
    
    def list_templates(self) -> List[str]:
        """List all available templates."""
        if not os.path.exists(self.templates_dir):
            return []
        
        templates = []
        for f in os.listdir(self.templates_dir):
            if f.endswith('.docx') and not f.startswith('~'):
                templates.append(f)
        return templates
